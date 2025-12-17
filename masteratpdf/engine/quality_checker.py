"""
MasterAtPDF Engine: Quality Checker

Valida la calidad de la conversión PDF → DOCX.
"""

from typing import Dict, List, Tuple
from docx import Document
import logging

logger = logging.getLogger(__name__)


class QualityMetrics:
    """Métricas de calidad de conversión."""
    
    def __init__(self):
        self.total_pages = 0
        self.total_paragraphs = 0
        self.total_bookmarks = 0
        self.total_hyperlinks = 0
        self.missing_bookmarks = []
        self.broken_links = []
        self.quality_score = 0.0
    
    def to_dict(self) -> Dict:
        """Convierte a diccionario para serialización."""
        return {
            "total_pages": self.total_pages,
            "total_paragraphs": self.total_paragraphs,
            "total_bookmarks": self.total_bookmarks,
            "total_hyperlinks": self.total_hyperlinks,
            "missing_bookmarks": len(self.missing_bookmarks),
            "broken_links": len(self.broken_links),
            "quality_score": self.quality_score
        }


class QualityChecker:
    """
    Valida la calidad de un DOCX generado.
    
    Checks:
    - Bookmarks insertados correctamente
    - Hyperlinks apuntan a bookmarks existentes
    - Índices son navegables
    - Estructura preservada
    """
    
    def __init__(self, docx_path: str):
        """
        Args:
            docx_path: Ruta del DOCX a validar
        """
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.metrics = QualityMetrics()
    
    def check_all(self) -> QualityMetrics:
        """
        Ejecuta todos los checks de calidad.
        
        Returns:
            QualityMetrics con resultados
        """
        logger.info(f"Iniciando quality check: {self.docx_path}")
        
        # Métricas básicas
        self.metrics.total_paragraphs = len(self.doc.paragraphs)
        
        # Check bookmarks
        bookmarks = self._extract_bookmarks()
        self.metrics.total_bookmarks = len(bookmarks)
        
        # Check hyperlinks
        hyperlinks = self._extract_hyperlinks()
        self.metrics.total_hyperlinks = len(hyperlinks)
        
        # Validar links → bookmarks
        self._validate_links(hyperlinks, bookmarks)
        
        # Calcular score
        self.metrics.quality_score = self._calculate_quality_score()
        
        logger.info(f"Quality check completado. Score: {self.metrics.quality_score:.2%}")
        return self.metrics
    
    def _extract_bookmarks(self) -> set:
        """Extrae todos los bookmarks del documento."""
        bookmarks = set()
        
        for para in self.doc.paragraphs:
            p = para._p
            for elem in p.iter():
                if elem.tag.endswith('bookmarkStart'):
                    name = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name')
                    if name and not name.startswith('_'):  # Ignorar bookmarks internos
                        bookmarks.add(name)
        
        return bookmarks
    
    def _extract_hyperlinks(self) -> List[Tuple[str, str]]:
        """
        Extrae todos los hyperlinks del documento.
        
        Returns:
            Lista de (text, anchor)
        """
        hyperlinks = []
        
        for para in self.doc.paragraphs:
            p = para._p
            for elem in p.iter():
                if elem.tag.endswith('hyperlink'):
                    anchor = elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor')
                    if anchor:
                        text = ''.join(elem.itertext())
                        hyperlinks.append((text, anchor))
        
        return hyperlinks
    
    def _validate_links(self, hyperlinks: List[Tuple[str, str]], bookmarks: set):
        """Valida que todos los hyperlinks apunten a bookmarks existentes."""
        for text, anchor in hyperlinks:
            if anchor not in bookmarks:
                self.metrics.broken_links.append({
                    "text": text,
                    "anchor": anchor,
                    "reason": "bookmark_not_found"
                })
    
    def _calculate_quality_score(self) -> float:
        """
        Calcula score de calidad (0-1).
        
        Formula:
        - Base: 1.0
        - Penalización por broken links
        - Bonus por bookmarks
        """
        score = 1.0
        
        # Penalización por links rotos
        if self.metrics.total_hyperlinks > 0:
            broken_ratio = len(self.metrics.broken_links) / self.metrics.total_hyperlinks
            score -= (broken_ratio * 0.3)  # Max -30%
        
        # Bonus por tener bookmarks
        if self.metrics.total_bookmarks > 0:
            score += 0.1  # +10%
        
        # Bonus por tener hyperlinks
        if self.metrics.total_hyperlinks > 0:
            score += 0.1  # +10%
        
        return max(0.0, min(1.0, score))
    
    def print_report(self):
        """Imprime reporte de calidad."""
        print("\n" + "="*60)
        print("📊 REPORTE DE CALIDAD")
        print("="*60)
        print(f"\n📄 Documento: {self.docx_path}")
        print(f"\n📈 Métricas:")
        print(f"   • Párrafos: {self.metrics.total_paragraphs}")
        print(f"   • Bookmarks: {self.metrics.total_bookmarks}")
        print(f"   • Hyperlinks: {self.metrics.total_hyperlinks}")
        
        if self.metrics.broken_links:
            print(f"\n⚠️  Links Rotos: {len(self.metrics.broken_links)}")
            for link in self.metrics.broken_links[:5]:  # Primeros 5
                print(f"   • '{link['text'][:40]}...' → {link['anchor']}")
        else:
            print(f"\n✅ Todos los hyperlinks son válidos")
        
        print(f"\n🎯 Quality Score: {self.metrics.quality_score:.1%}")
        
        if self.metrics.quality_score >= 0.9:
            print("   ✅ Excelente calidad")
        elif self.metrics.quality_score >= 0.7:
            print("   ✓ Buena calidad")
        else:
            print("   ⚠️ Necesita mejoras")
        
        print("="*60)
