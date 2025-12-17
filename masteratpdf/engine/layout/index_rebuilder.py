"""
MasterAtPDF: Index Rebuilder

Reconstruye el índice en DOCX con formato EXACTO del PDF.
Usa tablas invisibles para control preciso del layout.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.shared import qn
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls


class IndexRebuilder:
    """
    Reconstruye índice con layout pixel-perfect.
    
    Usa tablas invisibles para control total del posicionamiento.
    """
    
    def __init__(self, doc: Document = None):
        self.doc = doc if doc else Document()
    
    def rebuild_index_from_layout(self, entries: list, layout_data: dict):
        """
        Reconstruye el índice usando la data de layout extraída.
        
        Args:
            entries: Lista de entradas parseadas
            layout_data: Data del LayoutAnalyzer
        """
        print("\n🔨 Reconstruyendo índice con formato preciso...")
        
        # Título del índice
        self._add_index_title()
        
        # Agregar espacio
        self.doc.add_paragraph()
        
        # Usar tabla para layout preciso
        self._rebuild_with_table(entries)
        
        print(f"  ✓ {len(entries)} entradas reconstruidas")
    
    def _add_index_title(self):
        """Agrega el título 'Índice' con formato."""
        title = self.doc.add_paragraph()
        run = title.add_run("Índice")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _rebuild_with_table(self, entries: list):
        """
        Reconstruye índice usando PÁRRAFOS (no tablas).
        
        Las tablas no renderizan leader dots correctamente.
        Usamos párrafos con tab stops nativos.
        """
        # Filtrar entradas válidas
        valid_entries = [e for e in entries if len(e.get('text', '').strip()) > 5]
        
        if not valid_entries:
            print("  ⚠️ No hay entradas válidas")
            return
        
        print(f"  → Creando {len(valid_entries)} entradas con puntos guía...")
        
        # Crear cada entrada como un párrafo separado
        for entry in valid_entries:
            # Limpiar texto
            import re
            text = entry['text']
            # Remover número de página del final
            text = re.sub(r'\.{2,}\s*\d+\s*$', '', text)
            text = re.sub(r'\s+\d+\s*$', '', text)
            text = text.strip()
            
            if not text or len(text) < 3:
                continue
            
            # Crear párrafo
            para = self.doc.add_paragraph()
            
            # Configurar indentación según nivel
            level = entry.get('level', 1)
            indent_inches = (level - 1) * 0.25
            para.paragraph_format.left_indent = Inches(indent_inches)
            
            # CRÍTICO: Tab stop con LEADER DOTS
            # Debe estar ANTES de agregar el texto
            tab_stops = para.paragraph_format.tab_stops
            tab_stop_position = Inches(6.0)  # Posición del número de página
            tab_stops.add_tab_stop(
                tab_stop_position,
                WD_TAB_ALIGNMENT.RIGHT,
                WD_TAB_LEADER.DOTS  # ← Esto crea los puntos guía
            )
            
            # Agregar texto del título
            run = para.add_run(text)
            run.font.name = self._map_pdf_font(entry.get('font_name', 'Times'))
            run.font.size = Pt(entry.get('font_size', 12))
            if entry.get('is_bold'):
                run.font.bold = True
            if entry.get('is_italic'):
                run.font.italic = True
            
            # Tab (esto activa los puntos guía)
            para.add_run('\t')
            
            # Número de página
            page_num = entry.get('page', '')
            if page_num and page_num > 0:
                page_run = para.add_run(str(page_num))
                page_run.font.name = self._map_pdf_font(entry.get('font_name', 'Times'))
                page_run.font.size = Pt(entry.get('font_size', 12))
    
    def _make_table_invisible(self, table):
        """Hace los bordes de la tabla invisibles."""
        tbl = table._element
        tblPr = tbl.tblPr
        
        # Crear o obtener tblBorders
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        
        # Hacer todos los bordes "none"
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
    
    def _map_pdf_font(self, pdf_font_name: str) -> str:
        """Mapea fuentes PDF a fuentes Word."""
        FONT_MAP = {
            'TimesNewRomanPSMT': 'Times New Roman',
            'TimesNewRomanPS-BoldMT': 'Times New Roman',
            'ArialMT': 'Arial',
            'Arial-BoldMT': 'Arial',
            'Helvetica': 'Arial',
            'Courier': 'Courier New',
        }
        
        # Buscar coincidencia parcial
        for pdf_key, word_font in FONT_MAP.items():
            if pdf_key in pdf_font_name:
                return word_font
        
        # Default
        return 'Times New Roman'
    
    def add_hyperlinks_to_index(self, bookmarks: dict):
        """
        Agrega hyperlinks a las entradas del índice.
        
        Ahora trabaja con párrafos en vez de tablas.
        
        Args:
            bookmarks: {nombre_bookmark: para_idx}
        """
        print("\n🔗 Agregando hyperlinks al índice...")
        
        links_added = 0
        
        # Iterar sobre los primeros ~100 párrafos (donde está el índice)
        for para in self.doc.paragraphs[:100]:
            text = para.text.strip()
            
            if not text or len(text) < 5:
                continue
            
            # Skip título "Índice"
            if text == "Índice":
                continue
            
            # Buscar bookmark correspondiente
            bookmark_name = self._find_bookmark_for_text(text, bookmarks)
            
            if bookmark_name:
                # Convertir en hyperlink
                self._convert_paragraph_to_hyperlink_preserve_format(para, bookmark_name)
                links_added += 1
        
        print(f"  ✓ {links_added} hyperlinks agregados")
    
    def _find_bookmark_for_text(self, text: str, bookmarks: dict) -> str:
        """Encuentra el bookmark correspondiente a un texto."""
        import re
        
        # CAPITULO
        if 'CAPITULO' in text.upper():
            match = re.search(r'CAPITULO\s+([IVXLC]+)', text, re.IGNORECASE)
            if match:
                bookmark_name = f"capitulo_{match.group(1)}"
                if bookmark_name in bookmarks:
                    return bookmark_name
        
        # Sección numérica
        if re.match(r'^\d+\.\d+', text):
            match = re.match(r'^(\d+\.\d+(?:\.\d+)?)', text)
            if match:
                section_num = match.group(1)
                bookmark_name = f"section_{section_num.replace('.', '_')}"
                if bookmark_name in bookmarks:
                    return bookmark_name
        
        # INTRODUCCIÓN, CONCLUSIONES, etc.
        keywords = {
            'INTRODUCCIÓN': 'introduccion',
            'CONCLUSIONES': 'conclusiones',
            'RECOMENDACIONES': 'recomendaciones',
            'REFERENCIAS': 'referencias',
        }
        
        for keyword, bookmark_base in keywords.items():
            if keyword in text.upper():
                if bookmark_base in bookmarks:
                    return bookmark_base
        
        return None
    
    def _convert_paragraph_to_hyperlink(self, para, text, bookmark_name):
        """Convierte un párrafo en hyperlink."""
        # Guardar formato
        run_props = []
        for run in para.runs:
            run_props.append({
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.font.bold,
                'italic': run.font.italic
            })
        
        # Limpiar runs
        p = para._p
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        
        # Crear hyperlink
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Aplicar formato original (sin cambiar color/underline)
        if run_props:
            props = run_props[0]
            if props.get('bold'):
                b = OxmlElement('w:b')
                rPr.append(b)
            if props.get('italic'):
                i = OxmlElement('w:i')
                rPr.append(i)
        
        run.append(rPr)
        
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        run.append(t)
        
        hyperlink.append(run)
        p.append(hyperlink)
        d e f   _ c o n v e r t _ p a r a g r a p h _ t o _ h y p e r l i n k _ p r e s e r v e _ f o r m a t ( s e l f ,   p a r a ,   b o o k m a r k _ n a m e ) :  
                 " " "  
                 C o n v i e r t e   p � � r r a f o   e n   h y p e r l i n k   P R E S E R V A N D O   t a b   s t o p s   y   f o r m a t o .  
                  
                 E s t o   e s   c r � � t i c o   p a r a   m a n t e n e r   l o s   p u n t o s   g u � � a .  
                 " " "  
                 f r o m   d o c x . o x m l   i m p o r t   O x m l E l e m e n t  
                 f r o m   d o c x . o x m l . s h a r e d   i m p o r t   q n  
                  
                 #   E x t r a e r   e l   t e x t o   a n t e s   d e l   t a b  
                 f u l l _ t e x t   =   p a r a . t e x t  
                 p a r t s   =   f u l l _ t e x t . s p l i t ( ' \ t ' )  
                 i f   l e n ( p a r t s )   <   2 :  
                         r e t u r n     #   N o   h a y   t a b ,   n o   h a c e r   n a d a  
                  
                 t i t l e _ t e x t   =   p a r t s [ 0 ] . s t r i p ( )  
                 p a g e _ t e x t   =   p a r t s [ 1 ] . s t r i p ( )  
                  
                 #   G u a r d a r   f o r m a t o   d e l   t a b   s t o p  
                 t a b _ s t o p s   =   l i s t ( p a r a . p a r a g r a p h _ f o r m a t . t a b _ s t o p s )  
                 i n d e n t   =   p a r a . p a r a g r a p h _ f o r m a t . l e f t _ i n d e n t  
                  
                 #   L i m p i a r   p � � r r a f o  
                 p   =   p a r a . _ p  
                 f o r   r u n   i n   l i s t ( p a r a . r u n s ) :  
                         r u n . _ e l e m e n t . g e t p a r e n t ( ) . r e m o v e ( r u n . _ e l e m e n t )  
                  
                 #   R e c r e a r   c o n   h y p e r l i n k  
                 h y p e r l i n k   =   O x m l E l e m e n t ( ' w : h y p e r l i n k ' )  
                 h y p e r l i n k . s e t ( q n ( ' w : a n c h o r ' ) ,   b o o k m a r k _ n a m e )  
                  
                 #   R u n   c o n   t � � t u l o   ( d e n t r o   d e l   h y p e r l i n k )  
                 r u n 1   =   O x m l E l e m e n t ( ' w : r ' )  
                 r P r 1   =   O x m l E l e m e n t ( ' w : r P r ' )  
                  
                 #   C o l o r   a z u l   +   u n d e r l i n e   p a r a   h y p e r l i n k  
                 u   =   O x m l E l e m e n t ( ' w : u ' )  
                 u . s e t ( q n ( ' w : v a l ' ) ,   ' s i n g l e ' )  
                 c o l o r   =   O x m l E l e m e n t ( ' w : c o l o r ' )  
                 c o l o r . s e t ( q n ( ' w : v a l ' ) ,   ' 0 0 0 0 F F ' )  
                  
                 r P r 1 . a p p e n d ( u )  
                 r P r 1 . a p p e n d ( c o l o r )  
                 r u n 1 . a p p e n d ( r P r 1 )  
                  
                 t 1   =   O x m l E l e m e n t ( ' w : t ' )  
                 t 1 . s e t ( q n ( ' x m l : s p a c e ' ) ,   ' p r e s e r v e ' )  
                 t 1 . t e x t   =   t i t l e _ t e x t  
                 r u n 1 . a p p e n d ( t 1 )  
                  
                 h y p e r l i n k . a p p e n d ( r u n 1 )  
                 p . a p p e n d ( h y p e r l i n k )  
                  
                 #   T a b   ( F U E R A   d e l   h y p e r l i n k )  
                 r u n _ t a b   =   O x m l E l e m e n t ( ' w : r ' )  
                 t a b   =   O x m l E l e m e n t ( ' w : t a b ' )  
                 r u n _ t a b . a p p e n d ( t a b )  
                 p . a p p e n d ( r u n _ t a b )  
                  
                 #   N � � m e r o   d e   p � � g i n a   ( F U E R A   d e l   h y p e r l i n k )  
                 r u n 2   =   O x m l E l e m e n t ( ' w : r ' )  
                 t 2   =   O x m l E l e m e n t ( ' w : t ' )  
                 t 2 . t e x t   =   p a g e _ t e x t  
                 r u n 2 . a p p e n d ( t 2 )  
                 p . a p p e n d ( r u n 2 )  
                  
                 #   R e s t a u r a r   t a b   s t o p s  
                 p a r a . p a r a g r a p h _ f o r m a t . l e f t _ i n d e n t   =   i n d e n t  
                 f o r   t a b _ s t o p   i n   t a b _ s t o p s :  
                         p a r a . p a r a g r a p h _ f o r m a t . t a b _ s t o p s . a d d _ t a b _ s t o p (  
                                 t a b _ s t o p . p o s i t i o n ,  
                                 t a b _ s t o p . a l i g n m e n t ,  
                                 t a b _ s t o p . l e a d e r  
                         )  
 