"""
MasterAtPDF Engine: Hybrid Converter

NUEVA ESTRATEGIA:
1. Usar pdf2docx para conversión visual (preserva orden)
2. Analizar PDF para extraer estructura
3. Hacer navegable el contenido EXISTENTE
"""

import logging
import tempfile
import os
from pathlib import Path
from docx import Document
from pdf2docx import Converter

from .pdf_reader import PDFReader
from .structure_analyzer import StructureAnalyzer
from .docx_writer import DOCXWriter

logger = logging.getLogger(__name__)


class HybridConverter:
    """
    Conversor híbrido que combina lo mejor de ambos mundos:
    - pdf2docx: Conversión visual precisa
    - Motor nativo: Bookmarks y navegación
    """
    
    def __init__(self, verbose: bool = True):
        self.verbose = verbose
    
    def convert(self, pdf_path: str, output_path: str) -> bool:
        """
        Conversión híbrida optimizada.
        
        Args:
            pdf_path: PDF de entrada
            output_path: DOCX de salida
            
        Returns:
            True si exitoso
        """
        try:
            self._log("[INFO] 🚀 Conversión Híbrida (pdf2docx + Motor Nativo)")
            
            # ============================================
            # FASE 1: Conversión Visual con pdf2docx
            # ============================================
            self._log("[INFO] Paso 1/4: Conversión visual (pdf2docx)...")
            
            temp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            temp.close()
            
            cv = Converter(pdf_path)
            cv.convert(temp.name, start=0, end=None)
            cv.close()
            
            doc = Document(temp.name)
            self._log(f"  ✓ {len(doc.paragraphs)} párrafos convertidos")
            
            # ============================================
            # FASE 2: Analizar PDF para estructura
            # ============================================
            self._log("[INFO] Paso 2/4: Analizando estructura PDF...")
            
            with PDFReader(pdf_path) as reader:
                pdf_data = reader.read()
            
            analyzer = StructureAnalyzer(pdf_data)
            structure = analyzer.analyze()
            
            self._log(f"  ✓ {len(structure['figures'])} figuras")
            self._log(f"  ✓ {len(structure['tables'])} tablas")
            self._log(f"  ✓ {len(structure['headings'])} títulos")
            
            # ============================================
            # FASE 3: Insertar Bookmarks en DOCX
            # ============================================
            self._log("[INFO] Paso 3/4: Insertando bookmarks...")
            
            writer = DOCXWriter()
            writer.doc = doc  # Usar el documento ya convertido
            
            bookmarks_added = self._insert_bookmarks(writer, structure)
            self._log(f"  ✓ {bookmarks_added} bookmarks insertados")
            
            # ============================================
            # FASE 4: Hacer Índices Navegables
            # ============================================
            self._log("[INFO] Paso 4/4: Haciendo índices navegables...")
            
            links_created = self._make_indices_navigable(writer, structure)
            self._log(f"  ✓ {links_created} hyperlinks creados")
            
            # Guardar
            doc.save(output_path)
            os.unlink(temp.name)
            
            self._log(f"[SUCCESS] ✅ Conversión completada: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error: {e}", exc_info=True)
            return False
    
    def _insert_bookmarks(self, writer: DOCXWriter, structure: dict) -> int:
        """Inserta bookmarks en ubicaciones exactas."""
        count = 0
        
        # Figuras
        for fig in structure['figures']:
            if fig.number:
                para_idx = writer.find_paragraph_by_text(fig.content, min_similarity=0.5)
                if para_idx >= 0:
                    writer.insert_bookmark(
                        writer.doc.paragraphs[para_idx],
                        f"figura_{fig.number}"
                    )
                    count += 1
        
        # Tablas
        for table in structure['tables']:
            if table.number:
                para_idx = writer.find_paragraph_by_text(table.content, min_similarity=0.5)
                if para_idx >= 0:
                    writer.insert_bookmark(
                        writer.doc.paragraphs[para_idx],
                        f"tabla_{table.number}"
                    )
                    count += 1
        
        # Headings
        for heading in structure['headings']:
            para_idx = writer.find_paragraph_by_text(heading.content, min_similarity=0.6)
            if para_idx >= 0:
                # Generar bookmark name según tipo
                if 'CAPITULO' in heading.content.upper():
                    import re
                    match = re.search(r'CAPITULO\s+([IVXLC]+)', heading.content, re.IGNORECASE)
                    if match:
                        bookmark_name = f"capitulo_{match.group(1)}"
                    else:
                        bookmark_name = f"heading_{heading.page_number}"
                elif re.match(r'^\d+\.\d+', heading.content):
                    # Sección numérica
                    section_num = re.match(r'^(\d+\.\d+(?:\.\d+)?)', heading.content).group(1)
                    bookmark_name = f"section_{section_num.replace('.', '_')}"
                else:
                    bookmark_name = f"heading_{heading.page_number}"
                
                writer.insert_bookmark(writer.doc.paragraphs[para_idx], bookmark_name)
                count += 1
        
        return count
    
    def _make_indices_navigable(self, writer: DOCXWriter, structure: dict) -> int:
        """Convierte entradas de índice en hyperlinks."""
        count = 0
        
        for index_section in structure['indices']:
            if index_section.type == 'figures':
                for entry_text, page_num in index_section.entries:
                    para_idx = writer.find_paragraph_by_text(entry_text, min_similarity=0.4)
                    if para_idx >= 0:
                        import re
                        match = re.search(r'\d+', entry_text)
                        if match:
                            fig_num = match.group()
                            bookmark = f"figura_{fig_num}"
                            writer.create_hyperlink(
                                writer.doc.paragraphs[para_idx],
                                writer.doc.paragraphs[para_idx].text,
                                bookmark
                            )
                            count += 1
            
            elif index_section.type == 'tables':
                for entry_text, page_num in index_section.entries:
                    para_idx = writer.find_paragraph_by_text(entry_text, min_similarity=0.4)
                    if para_idx >= 0:
                        import re
                        match = re.search(r'\d+', entry_text)
                        if match:
                            table_num = match.group()
                            bookmark = f"tabla_{table_num}"
                            writer.create_hyperlink(
                                writer.doc.paragraphs[para_idx],
                                writer.doc.paragraphs[para_idx].text,
                                bookmark
                            )
                            count += 1
            
            elif index_section.type == 'general':
                # Índice general - más complejo
                for entry_text, page_num in index_section.entries:
                    para_idx = writer.find_paragraph_by_text(entry_text, min_similarity=0.5)
                    if para_idx >= 0:
                        # Determinar bookmark basado en el texto
                        import re
                        if 'CAPITULO' in entry_text.upper():
                            match = re.search(r'CAPITULO\s+([IVXLC]+)', entry_text, re.IGNORECASE)
                            if match:
                                bookmark = f"capitulo_{match.group(1)}"
                                writer.create_hyperlink(
                                    writer.doc.paragraphs[para_idx],
                                    writer.doc.paragraphs[para_idx].text,
                                    bookmark
                                )
                                count += 1
                        elif re.match(r'^\d+\.\d+', entry_text):
                            section_match = re.match(r'^(\d+\.\d+(?:\.\d+)?)', entry_text)
                            if section_match:
                                section_num = section_match.group(1)
                                bookmark = f"section_{section_num.replace('.', '_')}"
                                writer.create_hyperlink(
                                    writer.doc.paragraphs[para_idx],
                                    writer.doc.paragraphs[para_idx].text,
                                    bookmark
                                )
                                count += 1
        
        return count
    
    def _log(self, message: str):
        """Log condicional."""
        if self.verbose:
            print(message)
