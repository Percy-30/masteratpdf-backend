"""
MasterAtPDF: Table Writer

Escribe tablas en DOCX preservando:
- Estructura de filas/columnas
- Celdas merged
- Formato interno (bold, italic, fonts)
- Bordes
- Alineación
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from docx.table import _Cell
from typing import Optional

from table_extractor import Table, CellContent, CellAlignment


class TableWriter:
    """Escribe tablas en DOCX con formato completo."""
    
    def __init__(self, doc: Optional[Document] = None):
        self.doc = doc if doc else Document()
    
    def write_table(self, table: Table) -> None:
        """
        Escribe tabla en documento DOCX.
        
        Args:
            table: Tabla extraída
        """
        
        if table.rows == 0 or table.cols == 0:
            return
        
        # Crear tabla en DOCX
        docx_table = self.doc.add_table(rows=table.rows, cols=table.cols)
        
        # Aplicar estilo básico
        docx_table.style = 'Table Grid'
        
        # Llenar celdas con contenido
        for i, row_data in enumerate(table.data):
            for j, cell_content in enumerate(row_data):
                cell = docx_table.rows[i].cells[j]
                self._write_cell_content(cell, cell_content)
        
        # Aplicar merged cells
        for merge in table.merged_cells:
            try:
                start_cell = docx_table.rows[merge.start_row].cells[merge.start_col]
                end_cell = docx_table.rows[merge.end_row].cells[merge.end_col]
                start_cell.merge(end_cell)
            except Exception as e:
                # Si falla merge, continuar
                pass
        
        # Agregar espacio después de tabla
        self.doc.add_paragraph()
    
    def _write_cell_content(self, cell: _Cell, content: CellContent) -> None:
        """
        Escribe contenido en una celda con formato.
        
        Args:
            cell: Celda de DOCX
            content: Contenido a escribir
        """
        
        # Limpiar párrafo existente
        para = cell.paragraphs[0]
        para.clear()
        
        # Agregar texto con formato
        run = para.add_run(content.text)
        
        # Aplicar formato de fuente
        run.font.name = content.font_name
        run.font.size = Pt(content.font_size)
        
        if content.is_bold:
            run.font.bold = True
        
        if content.is_italic:
            run.font.italic = True
        
        # Aplicar alineación
        if content.alignment == CellAlignment.CENTER:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif content.alignment == CellAlignment.RIGHT:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif content.alignment == CellAlignment.JUSTIFY:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def save(self, output_path: str) -> None:
        """Guarda documento."""
        self.doc.save(output_path)


if __name__ == "__main__":
    # Test completo: Detectar + Extraer + Escribir
    import fitz
    from table_detector import TableDetector
    from table_extractor import TableExtractor
    
    pdf_path = "d:/PROYECTOS/PYTHON/masteratpdf-backend/tesis.pdf"
    output_path = "d:/PROYECTOS/PYTHON/masteratpdf-backend/test_tables.docx"
    
    print("🔍 Pipeline completo de tablas...")
    
    # 1. Detectar
    print("\n[1/3] Detectando tablas...")
    detector = TableDetector(dpi=200)
    pdf_doc = fitz.open(pdf_path)
    
    # Test con página 52 (tiene Tabla 1)
    page = pdf_doc[51]
    table_regions = detector.detect_tables_in_page(page)
    print(f"  ✓ {len(table_regions)} tablas detectadas")
    
    # 2. Extraer
    print("\n[2/3] Extrayendo contenido...")
    extractor = TableExtractor()
    tables = []
    
    for region in table_regions[:5]:  # Primeras 5 tablas
        table = extractor.extract_table(page, region)
        if table:
            tables.append(table)
            print(f"  ✓ Tabla {table.rows}x{table.cols}")
    
    # 3. Escribir
    print("\n[3/3] Escribiendo a DOCX...")
    writer = TableWriter()
    
    writer.doc.add_heading("Tablas Extraídas de tesis.pdf", level=1)
    writer.doc.add_paragraph(f"Página 52 - {len(tables)} tablas")
    writer.doc.add_paragraph()
    
    for i, table in enumerate(tables, 1):
        writer.doc.add_heading(f"Tabla {i}", level=2)
        writer.write_table(table)
    
    writer.save(output_path)
    
    pdf_doc.close()
    
    print(f"\n✅ COMPLETADO: {output_path}")
    print(f"   {len(tables)} tablas escritas con formato completo")
