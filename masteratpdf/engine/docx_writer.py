"""
MasterAtPDF Engine: DOCX Writer

Construye DOCX desde cero con control total de posicionamiento.
NO usa pdf2docx - construcción directa.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Tuple
import logging

from .coordinate_mapper import CoordinateMapper, PDFRect, DOCXPosition
from .pdf_reader import Block

logger = logging.getLogger(__name__)


class DOCXWriter:
    """
    Construye DOCX directamente desde estructura PDF.
    
    Ventajas sobre pdf2docx:
    - Control total de posicionamiento
    - Bookmarks en ubicaciones exactas
    - Paginación preservada
    - Links internos precisos
    """
    
    def __init__(self):
        self.doc = Document()
        self.current_page = 1
        self.bookmark_counter = 0
    
    def write(self, pdf_structure: Dict, output_path: str):
        """
        Escribe DOCX desde estructura PDF.
        
        Args:
            pdf_structure: Output de PDFReader.read()
            output_path: Ruta del DOCX de salida
        """
        logger.info(f"Iniciando construcción de DOCX: {output_path}")
        
        pages = pdf_structure['pages']
        
        for page_data in pages:
            self._add_page(page_data)
        
        # Guardar
        self.doc.save(output_path)
        logger.info(f"✅ DOCX guardado: {output_path}")
    
    def _add_page(self, page_data: Dict):
        """Agrega una página completa al documento."""
        page_num = page_data['number']
        logger.debug(f"Procesando página {page_num}")
        
        # Crear mapper de coordenadas para esta página
        page_width, page_height = page_data['size']
        mapper = CoordinateMapper(page_width, page_height)
        
        # CRÍTICO: Ordenar bloques TOP-TO-BOTTOM (Y decrece en PDF = top)
        # En PDF, Y=0 está abajo, Y=max está arriba
        # Queremos procesar de arriba (Y alto) a abajo (Y bajo)
        # Ordenar por Y1 descendente (primero los de arriba)
        blocks = sorted(page_data['blocks'], key=lambda b: (-b.y1, b.x0))
        
        for block in blocks:
            if block.type == 'text':
                self._add_text_block(block, mapper)
            elif block.type == 'image':
                self._add_image_block(block, mapper)
        
        # Salto de página (excepto última)
        self.doc.add_page_break()
        self.current_page += 1
    
    def _add_text_block(self, block: Block, mapper: CoordinateMapper):
        """Agrega un bloque de texto con formato."""
        lines = block.content.get('lines', [])
        
        for line in lines:
            # Crear párrafo
            para = self.doc.add_paragraph()
            
            # Procesar spans (runs) de la línea
            for span in line.get('spans', []):
                text = span.get('text', '')
                if not text.strip():
                    continue
                
                run = para.add_run(text)
                
                # Aplicar estilo del span
                self._apply_span_style(run, span)
        
    def _apply_span_style(self, run, span: Dict):
        """Aplica estilos de un span PDF a un run DOCX."""
        # Font size
        size = span.get('size', 12)
        run.font.size = Pt(size)
        
        # Font name
        font_name = span.get('font', 'Times New Roman')
        run.font.name = font_name
        
        # Bold/Italic
        flags = span.get('flags', 0)
        if flags & 16:  # Bold
            run.font.bold = True
        if flags & 2:   # Italic
            run.font.italic = True
        
        # Color
        color = span.get('color', None)
        if color and isinstance(color, int):
            # Convert int color to RGB
            r = (color >> 16) & 0xFF
            g = (color >> 8) & 0xFF
            b = color & 0xFF
            run.font.color.rgb = RGBColor(r, g, b)
    
    def _add_image_block(self, block: Block, mapper: CoordinateMapper):
        """Agrega una imagen al documento."""
        # TODO: Implementar extracción y colocación de imágenes
        pass
    
    def insert_bookmark(self, paragraph, bookmark_name: str) -> bool:
        """
        Inserta un bookmark en un párrafo específico.
        
        Args:
            paragraph: Párrafo de python-docx
            bookmark_name: Nombre del bookmark
            
        Returns:
            True si se insertó correctamente
        """
        try:
            bookmark_id = str(self.bookmark_counter)
            self.bookmark_counter += 1
            
            p = paragraph._p
            
            # Bookmark start
            start = OxmlElement('w:bookmarkStart')
            start.set(qn('w:id'), bookmark_id)
            start.set(qn('w:name'), bookmark_name)
            
            # Bookmark end
            end = OxmlElement('w:bookmarkEnd')
            end.set(qn('w:id'), bookmark_id)
            
            # Insertar
            p.insert(0, start)
            p.append(end)
            
            logger.debug(f"Bookmark insertado: {bookmark_name}")
            return True
            
        except Exception as e:
            logger.error(f"Error al insertar bookmark {bookmark_name}: {e}")
            return False
    
    def create_hyperlink(self, paragraph, text: str, bookmark_name: str):
        """
        Crea un hyperlink en un párrafo que apunta a un bookmark.
        
        Args:
            paragraph: Párrafo de python-docx
            text: Texto del link
            bookmark_name: Nombre del bookmark destino
        """
        # Limpiar runs existentes
        p = paragraph._p
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
        
        # Crear hyperlink
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)
        
        # Crear run
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Estilo de hyperlink
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        
        rPr.append(u)
        rPr.append(color)
        run.append(rPr)
        
        # Texto
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        run.append(t)
        
        hyperlink.append(run)
        p.append(hyperlink)
    
    def find_paragraph_by_text(self, search_text: str, min_similarity: float = 0.6) -> int:
        """
        Encuentra un párrafo por texto con búsqueda optimizada.
        
        Args:
            search_text: Texto a buscar
            min_similarity: Similitud mínima (0-1)
            
        Returns:
            Índice del párrafo o -1 si no se encuentra
        """
        from .text_matcher import TextMatcher
        
        # Extraer textos de párrafos
        para_texts = [para.text for para in self.doc.paragraphs]
        
        # Usar TextMatcher mejorado
        best_idx = TextMatcher.find_best_match(search_text, para_texts, min_similarity)
        
        return best_idx if best_idx is not None else -1
