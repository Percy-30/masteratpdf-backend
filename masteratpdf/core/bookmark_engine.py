from docx import Document
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
import re
from typing import Optional

def find_paragraph_by_text(doc: Document, search_text: str, partial: bool = True, min_similarity: float = 0.6) -> Optional[int]:
    """
    Busca un párrafo que contenga el texto especificado con búsqueda flexible.
    """
    search_clean = re.sub(r'\s+', ' ', search_text.strip().lower())
    search_words = set(search_clean.split())
    
    best_match_idx = None
    best_similarity = 0
    
    for idx, paragraph in enumerate(doc.paragraphs):
        para_text = re.sub(r'\s+', ' ', paragraph.text.strip().lower())
        
        if partial:
            # Búsqueda por similitud de palabras
            para_words = set(para_text.split())
            if search_words and para_words:
                common_words = search_words & para_words
                similarity = len(common_words) / len(search_words)
                
                if similarity > best_similarity and similarity >= min_similarity:
                    best_similarity = similarity
                    best_match_idx = idx
        else:
            if para_text == search_clean:
                return idx
    
    return best_match_idx


def insert_bookmark_at_paragraph(paragraph, bookmark_name: str, bookmark_id: str = None):
    """
    Inserta un bookmark en un párrafo específico.
    """
    if bookmark_id is None:
        bookmark_id = str(abs(hash(bookmark_name)) % 100000)
    
    p = paragraph._p
    
    # Verificar si ya tiene bookmark con este nombre
    for elem in p.iter():
        if elem.tag == qn('w:bookmarkStart'):
            if elem.get(qn('w:name')) == bookmark_name:
                return  # Ya existe
    
    # Crear bookmark start
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), bookmark_id)
    start.set(qn('w:name'), bookmark_name)
    
    # Crear bookmark end
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bookmark_id)
    
    # Insertar al principio del párrafo
    p.insert(0, start)
    p.append(end)


def place_toc_bookmarks(doc: Document, toc_items: list):
    """
    Coloca bookmarks para TOC general.
    """
    for level, title, page in toc_items:
        bookmark_name = f"toc_{title[:30].replace(' ', '_')}"
        
        # Buscar por título
        para_idx = find_paragraph_by_text(doc, title, partial=True, min_similarity=0.5)
        
        if para_idx is not None:
            insert_bookmark_at_paragraph(doc.paragraphs[para_idx], bookmark_name)


def place_figure_bookmarks(doc: Document, figures: list):
    """
    Coloca bookmarks para cada figura con búsqueda mejorada.
    """
    for title, page in figures:
        # Extraer número de figura
        match = re.search(r'Figura\s+(\d+)', title, re.IGNORECASE)
        if not match:
            continue
            
        fig_num = match.group(1)
        bookmark_name = f"figura_{fig_num}"
        
        # Buscar múltiples variantes
        search_variants = [
            f"Figura {fig_num}",
            f"Figura{fig_num}",
            title[:50],  # Primeros 50 caracteres
        ]
        
        para_idx = None
        for variant in search_variants:
            para_idx = find_paragraph_by_text(doc, variant, partial=True, min_similarity=0.4)
            if para_idx is not None:
                break
        
        if para_idx is not None:
            insert_bookmark_at_paragraph(doc.paragraphs[para_idx], bookmark_name)


def place_table_bookmarks(doc: Document, tables: list):
    """
    Coloca bookmarks para cada tabla con búsqueda mejorada.
    """
    for title, page in tables:
        # Extraer número de tabla
        match = re.search(r'Tabla\s+(\d+)', title, re.IGNORECASE)
        if not match:
            continue
            
        table_num = match.group(1)
        bookmark_name = f"tabla_{table_num}"
        
        # Buscar múltiples variantes
        search_variants = [
            f"Tabla {table_num}",
            f"Tabla{table_num}",
            title[:50],
        ]
        
        para_idx = None
        for variant in search_variants:
            para_idx = find_paragraph_by_text(doc, variant, partial=True, min_similarity=0.4)
            if para_idx is not None:
                break
        
        if para_idx is not None:
            insert_bookmark_at_paragraph(doc.paragraphs[para_idx], bookmark_name)


def create_hyperlink(paragraph, text: str, bookmark: str):
    """
    Añade un hyperlink a un párrafo que apunta a un bookmark.
    """
    p = paragraph._p
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark)
    
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    
    rPr.append(u)
    rPr.append(color)
    run.append(rPr)
    
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    
    hyperlink.append(run)
    p.append(hyperlink)
